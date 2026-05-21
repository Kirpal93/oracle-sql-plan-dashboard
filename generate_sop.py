"""
Generates the Oracle SQL Plan Dashboard — Scope, Approach & SOP
as a formatted Word (.docx) document and then converts it to PDF.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os
from pathlib import Path

OUT_DIR  = Path(__file__).parent
DOCX_OUT = OUT_DIR / "Oracle_SQL_Plan_Dashboard_SOP.docx"
PDF_OUT  = OUT_DIR / "Oracle_SQL_Plan_Dashboard_SOP.pdf"

# ── colour palette ─────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x3F, 0x7A)   # headings
BLUE_MID   = RGBColor(0x2E, 0x6D, 0xC8)   # sub-headings
BLUE_LIGHT = RGBColor(0xD6, 0xE4, 0xF7)   # table header fill
GREY_LIGHT = RGBColor(0xF2, 0xF5, 0xFA)   # alternating row
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)
GREEN_DARK = RGBColor(0x1E, 0x6B, 0x3A)


# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a hex colour string (e.g. 'D6E4F7')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_cell_border(table):
    """Add thin borders to every cell in a table."""
    tbl  = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def style_header_row(row, bg_hex: str = "1F3F7A"):
    for cell in row.cells:
        set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold      = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None, alt_rows=True):
    """Create a styled table and return it."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_cell_border(table)

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, "1F3F7A")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = "F2F5FA" if (alt_rows and r_idx % 2 == 1) else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = BLACK

    # column widths
    if col_widths:
        for col_idx, width in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(width)

    return table


def add_heading(doc, text, level=1, color=None):
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.color.rgb = color or (BLUE_DARK if level == 1 else BLUE_MID)
    return para


def add_para(doc, text, bold=False, italic=False, size=10, color=None, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Inches(0.3)
    run  = para.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    run.font.color.rgb = color or BLACK
    return para


def add_code_block(doc, code_text):
    """Add a monospaced code block paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    # light grey shading via XML
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EEF2F7")
    pPr.append(shd)
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    run  = para.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    return para


def divider(doc):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DC8")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover / Title ──────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run  = title_para.add_run("Oracle SQL Plan Fluctuation Dashboard")
title_run.bold            = True
title_run.font.size       = Pt(22)
title_run.font.color.rgb  = BLUE_DARK

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run  = sub_para.add_run("Scope, Approach & Standard Operating Procedure (SOP)")
sub_run.font.size      = Pt(13)
sub_run.font.color.rgb = BLUE_MID
sub_run.italic         = True

meta_para = doc.add_paragraph()
meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run  = meta_para.add_run("Version 1.0  |  May 2026  |  Confidential")
meta_run.font.size      = Pt(9)
meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

divider(doc)
doc.add_paragraph()

# ── Section 1 — Executive Summary ─────────────────────────────────────────────
add_heading(doc, "1. Executive Summary", level=1)
add_para(doc,
    "The Oracle SQL Plan Fluctuation Dashboard is a Streamlit-based web application that "
    "provides a single-pane-of-glass view of SQL execution plan instability across one or more "
    "Oracle databases. It empowers DBAs and performance engineers to detect, diagnose, and "
    "remediate plan regressions without writing manual queries — and without requiring Oracle "
    "Enterprise Manager (OEM).",
    size=10)
doc.add_paragraph()

add_para(doc,
    "The dashboard connects to up to 50+ Oracle databases in parallel, surfaces all SQL IDs "
    "whose optimizer has chosen more than one execution plan, and presents historical trend data "
    "from AWR alongside automated tuning recommendations from Oracle's SQL Tuning Advisor.",
    size=10)
divider(doc)

# ── Section 2 — Scope ─────────────────────────────────────────────────────────
add_heading(doc, "2. Scope", level=1)

add_heading(doc, "2.1  In Scope", level=2)
in_scope = [
    ["Multi-Database Support",       "Query up to 50+ Oracle databases simultaneously in parallel (configurable thread pool)"],
    ["Plan Instability Detection",   "Identify SQL IDs with more than one PLAN_HASH_VALUE in the cursor cache (V$SQL)"],
    ["Module-Level Grouping",        "Group SQL IDs by Oracle MODULE / ACTION to attribute instability to application components"],
    ["Cross-DB Comparison",          "Side-by-side comparison of plan stability, execution counts, and elapsed time across all databases"],
    ["AWR Historical Analysis",      "Read DBA_HIST_SQLSTAT / DBA_HIST_SNAPSHOT to show per-snapshot performance and plan trends over time"],
    ["AWR Report Generation",        "Load profile, top SQL by elapsed time, and wait event analysis for a selected AWR snapshot range"],
    ["SQL Tuning Advisor",           "Run Oracle DBMS_SQLTUNE for any SQL ID and present structured findings (SQL Profile, Index, Statistics, Restructure)"],
    ["AI / LLM Q&A",                 "Natural-language question answering using a local Ollama/LLaMA3 model against the loaded data"],
    ["Data Export",                  "CSV export of module-level and SQL-level analysis results"],
]
add_table(doc, ["Feature", "Description"], in_scope, col_widths=[2.2, 4.4])
doc.add_paragraph()

add_heading(doc, "2.2  Out of Scope", level=2)
out_scope = [
    ["Database Changes",         "The dashboard is strictly read-only — no DDL/DML or configuration changes are made"],
    ["OEM Replacement",          "This is a targeted instability tool, not a full replacement for Oracle Enterprise Manager"],
    ["RAC / GoldenGate",         "RAC-level diagnostics and GoldenGate replication monitoring are not covered"],
    ["AWR Licensing",            "AWR-dependent tabs (Execution History, AWR Report) require an Oracle Diagnostics Pack license"],
    ["Tuning Pack Licensing",    "The SQL Tuning Advisor tab requires an Oracle Tuning Pack license"],
]
add_table(doc, ["Area", "Detail"], out_scope, col_widths=[2.2, 4.4])
divider(doc)

# ── Section 3 — Approach ──────────────────────────────────────────────────────
add_heading(doc, "3. Approach", level=1)

add_heading(doc, "3.1  Architecture Overview", level=2)
add_para(doc, "The solution follows a lightweight, agent-free architecture:", size=10)
arch_lines = [
    "Oracle Databases (1–50+)  →  Python oracledb driver (Thin mode — no Oracle Client install needed)",
    "Parallel query execution  →  Python ThreadPoolExecutor (20 threads by default)",
    "Data aggregation          →  Pandas DataFrames (in-memory, session-cached)",
    "Visualisation             →  Streamlit web app + Plotly interactive charts",
    "UI                        →  Dark-themed custom CSS, runs at http://localhost:8501",
]
add_code_block(doc, "\n".join(arch_lines))
doc.add_paragraph()

add_heading(doc, "3.2  Data Flow", level=2)
flow = [
    ["1", "Credentials Entry",       "DBA enters DB name, username, password, DSN in the sidebar Database Manager. Configs are saved to db_configs.json."],
    ["2", "Parallel Query",          "On 'Load Data', all selected databases are queried simultaneously using a thread pool. V$SQL is the primary source."],
    ["3", "Result Aggregation",      "Each database result is tagged with the database name and combined into a single Pandas DataFrame."],
    ["4", "Session Cache",           "Combined data is stored in Streamlit session state — no re-query occurs on every UI interaction."],
    ["5", "In-Memory Filtering",     "Sidebar filters (SQL ID, Schema, Module) are applied on the cached DataFrame without hitting the database again."],
    ["6", "AWR On-Demand",           "AWR/history data is fetched separately when the user selects a snapshot range in the relevant tabs."],
    ["7", "AI Context Assembly",     "For the Ask tab, a text summary of loaded data is assembled and sent to the local LLM for question answering."],
]
add_table(doc, ["Step", "Stage", "Detail"], flow, col_widths=[0.4, 1.6, 4.6])
doc.add_paragraph()

add_heading(doc, "3.3  Oracle Views Used", level=2)
views = [
    ["V$SQL",                  "Current cursor cache — plan hash values, execution stats, module/action"],
    ["V$SQL_PLAN",             "Execution plan steps for live (in-memory) cursors"],
    ["V$SQLAREA",              "Full SQL text retrieval (CLOB)"],
    ["DBA_HIST_SNAPSHOT",      "AWR snapshot catalogue — snap IDs and time intervals"],
    ["DBA_HIST_SQLSTAT",       "Per-snapshot SQL performance history (executions, elapsed time, etc.)"],
    ["DBA_HIST_SQL_PLAN",      "Historical execution plan steps stored in AWR"],
    ["DBA_HIST_SYSTEM_EVENT",  "Wait event history per AWR snapshot"],
    ["DBA_HIST_SYSSTAT",       "System-level load profile statistics per AWR snapshot"],
]
add_table(doc, ["Oracle View", "Purpose"], views, col_widths=[2.4, 4.2])
divider(doc)

# ── Section 4 — Dashboard Tabs ────────────────────────────────────────────────
add_heading(doc, "4. Dashboard Modules (Tabs)", level=1)
tabs = [
    ["Overview",               "KPI tiles (databases queried, total SQLs, fluctuating count, stable count, max plans). Pie chart (stable vs fluctuating). Bar chart (top 15 SQLs by plan count)."],
    ["Cross-DB Comparison",    "Side-by-side view of plan instability, execution totals, and avg elapsed time across all connected databases."],
    ["Module Drilldown",       "Groups SQL IDs by Oracle MODULE. KPI tiles per module, stacked bar chart (stable vs fluctuating per module), per-module drilldown with donut chart, health banner, and CSV export."],
    ["Plan Analysis",          "Select a SQL ID to view all plan variants, compare performance metrics (elapsed, CPU, buffer gets, disk reads) and view plan step trees side by side."],
    ["SQL Tuning Advisor",     "Run Oracle DBMS_SQLTUNE on any SQL ID. Displays structured findings: SQL Profile (blue), Index (amber), Statistics (purple), Restructure SQL (red). Supports accepting SQL Profile."],
    ["Execution History",      "AWR-based view. Select DB + snapshot range to see how a SQL's plan and performance evolved snapshot by snapshot. Line charts for avg elapsed, CPU, buffer gets."],
    ["AWR Report",             "Select snapshot range to generate: load profile (user calls, parses, physical reads/writes), top 50 SQLs by elapsed time with plan change indicator, top 20 wait events by class."],
    ["Ask the Dashboard",      "Natural-language Q&A. Type questions like 'Which database has the most plan fluctuations?' or 'List the top 5 unstable SQLs'. Answered from loaded data (rule-based + optional LLM)."],
]
add_table(doc, ["Tab", "Description"], tabs, col_widths=[2.0, 4.6])
divider(doc)

# ── Section 5 — Prerequisites ─────────────────────────────────────────────────
add_heading(doc, "5. Pre-Requisites", level=1)

add_heading(doc, "5.1  Oracle Database Grants", level=2)
add_para(doc, "The following grants must be applied by the customer DBA to the monitoring user before the dashboard can be used:", size=10)
doc.add_paragraph()

add_para(doc, "Core Dashboard (mandatory for all tabs):", bold=True, size=9.5)
add_code_block(doc,
"GRANT SELECT ON V_$SQL          TO <monitoring_user>;\n"
"GRANT SELECT ON V_$SQL_PLAN     TO <monitoring_user>;\n"
"GRANT SELECT ON V_$SQLAREA      TO <monitoring_user>;")

add_para(doc, "AWR / History Tabs (requires Oracle Diagnostics Pack license):", bold=True, size=9.5)
add_code_block(doc, "GRANT SELECT_CATALOG_ROLE       TO <monitoring_user>;")

add_para(doc, "SQL Tuning Advisor Tab (requires Oracle Tuning Pack license):", bold=True, size=9.5)
add_code_block(doc,
"GRANT ADVISOR                   TO <monitoring_user>;\n"
"GRANT EXECUTE ON DBMS_SQLTUNE   TO <monitoring_user>;")
doc.add_paragraph()

add_heading(doc, "5.2  Network Requirements", level=2)
net = [
    ["Listener Port",   "TCP port 1521 (or customer-configured port) must be reachable from the dashboard host to each Oracle host"],
    ["DSN Format",      "host:port/service_name  (e.g.  mydbserver:1521/ORCL)"],
    ["Firewall",        "Outbound TCP allowed from dashboard server to all Oracle DB servers on listener port"],
]
add_table(doc, ["Item", "Requirement"], net, col_widths=[1.8, 4.8])
doc.add_paragraph()

add_heading(doc, "5.3  Software Requirements", level=2)
sw = [
    ["Python",          "3.9 or higher"],
    ["streamlit",       "Web framework for the dashboard UI"],
    ["oracledb",        "Oracle database driver — Thin mode (no Oracle Instant Client needed)"],
    ["pandas",          "Data manipulation and in-memory filtering"],
    ["plotly",          "Interactive charts (pie, bar, line, radar)"],
    ["matplotlib",      "Supporting chart rendering"],
    ["python-dotenv",   "Loads .env configuration (Ollama URL, etc.)"],
    ["Ollama + LLaMA3", "Optional — required only for the AI Q&A tab. Must run locally on the dashboard host."],
]
add_table(doc, ["Package", "Purpose"], sw, col_widths=[1.8, 4.8])
divider(doc)

# ── Section 6 — SOP ───────────────────────────────────────────────────────────
add_heading(doc, "6. Standard Operating Procedure (SOP)", level=1)

# Step 1
add_heading(doc, "Step 1 — Installation (One-Time)", level=2)
add_para(doc, "Run the following command from the project folder to install all Python dependencies:", size=10)
add_code_block(doc, "python -m pip install -r requirements.txt")
doc.add_paragraph()

# Step 2
add_heading(doc, "Step 2 — Configure Environment (Optional — AI Q&A only)", level=2)
add_para(doc, "Create or edit the .env file in the project folder:", size=10)
add_code_block(doc, "OLLAMA_BASE_URL=http://localhost:11434")
doc.add_paragraph()

# Step 3
add_heading(doc, "Step 3 — Start the Dashboard", level=2)
add_para(doc, "Launch the Streamlit app from the project folder:", size=10)
add_code_block(doc, "python -m streamlit run oracle_sql_plan_dashboard.py")
add_para(doc, "The dashboard opens automatically in your browser at:  http://localhost:8501", size=10, italic=True)
doc.add_paragraph()

# Step 4
add_heading(doc, "Step 4 — Add Database Connections", level=2)
sop4 = [
    ["4a", "Open the sidebar",          "Click the '>' arrow on the left side if the sidebar is collapsed"],
    ["4b", "Open Database Manager",     "Locate the '🗄️ Database Manager' section at the top of the sidebar"],
    ["4c", "Add a new database",        "Click '➕ Add New Database', enter: Name (label), Username, Password, DSN (host:port/service)"],
    ["4d", "Test connection",           "Click 'Test Connection' to verify credentials and network access before saving"],
    ["4e", "Save",                       "Click 'Add Database' to persist the config to db_configs.json"],
    ["4f", "Repeat",                    "Repeat steps 4c–4e for each Oracle database to be monitored"],
]
add_table(doc, ["Sub-Step", "Action", "Detail"], sop4, col_widths=[0.7, 1.8, 4.1])
doc.add_paragraph()

# Step 5
add_heading(doc, "Step 5 — Load Data", level=2)
sop5 = [
    ["5a", "Select databases",          "Tick the checkboxes next to the databases you want to query in the sidebar"],
    ["5b", "Apply filters (optional)",  "Enter SQL ID, Schema, or Module filters in the sidebar if narrowing scope"],
    ["5c", "Load",                       "Click 'Load / Refresh Data' — all selected databases are queried in parallel"],
    ["5d", "Review errors",             "If any databases failed, connection errors are shown under the data load section"],
]
add_table(doc, ["Sub-Step", "Action", "Detail"], sop5, col_widths=[0.7, 1.8, 4.1])
doc.add_paragraph()

# Step 6
add_heading(doc, "Step 6 — Analyse Results by Tab", level=2)
sop6 = [
    ["Overview",           "Check the 5 KPI tiles for a quick health summary. Review pie and bar charts to assess instability spread."],
    ["Cross-DB Comparison","Identify which databases have the highest fluctuation counts. Use for prioritising remediation effort."],
    ["Module Drilldown",   "Identify which application module (e.g. ERP batch, web tier) is the primary source of instability. Export CSV to share with the development team."],
    ["Plan Analysis",      "Click a SQL ID in the summary table or enter it in the sidebar. Compare all plan variants. Note the fastest plan (lowest avg elapsed)."],
    ["SQL Tuning Advisor", "For fluctuating SQLs, run the Advisor. Review findings — accept a SQL Profile to pin the optimal plan."],
    ["Execution History",  "Select a DB, then choose a begin/end snapshot. Observe if the plan changed at a specific point in time (correlated to a stats refresh, code deploy, etc.)."],
    ["AWR Report",         "Select a snapshot range to see the overall database load profile and identify if the plan instability correlates with broader wait events or high parse rates."],
    ["Ask the Dashboard",  "Type questions in plain English — e.g. 'Which database has the most fluctuating SQLs?' or 'What is the worst performing module?'"],
]
add_table(doc, ["Tab", "Recommended Action"], sop6, col_widths=[2.0, 4.6])
doc.add_paragraph()

# Step 7
add_heading(doc, "Step 7 — Remediation Actions", level=2)
add_para(doc, "Based on analysis findings, apply the appropriate remediation:", size=10)
rem = [
    ["Plan Instability",         "Run SQL Tuning Advisor → accept SQL Profile to pin the optimal plan"],
    ["Missing Index",            "Forward the Advisor's index recommendation to the DBA team for review and implementation"],
    ["Stale Statistics",         "Schedule DBMS_STATS collection on the identified table(s): DBMS_STATS.GATHER_TABLE_STATS(...)"],
    ["Excessive Full Table Scan","Add appropriate index or rewrite the SQL to use selective predicates"],
    ["High Wait Events",         "Investigate the specific wait class from the AWR Report (e.g. I/O waits → storage review, locking → application code)"],
    ["High Hard Parse Rate",     "Review cursor sharing settings (CURSOR_SHARING parameter) and ensure bind variables are used in application SQLs"],
]
add_table(doc, ["Finding", "Remediation"], rem, col_widths=[2.2, 4.4])
divider(doc)

# ── Section 7 — Error Reference ───────────────────────────────────────────────
add_heading(doc, "7. Common Oracle Errors & Fixes", level=1)
errors = [
    ["ORA-00942", "Missing SELECT on V$SQL/V$SQL_PLAN/V$SQLAREA",          "GRANT SELECT ON V_$SQL TO <user>; (and V_$SQL_PLAN, V_$SQLAREA)"],
    ["ORA-01017", "Invalid username or password",                           "Verify credentials; check if the account is locked"],
    ["ORA-12541", "No listener / TNS no listener",                          "Verify DSN host:port/service — confirm listener is running"],
    ["ORA-12154", "TNS could not resolve service name",                     "Check the DSN format — use host:port/service_name (not a tnsnames alias)"],
    ["ORA-13773", "Missing cursor cache access for Advisor",                "GRANT SELECT_CATALOG_ROLE TO <user>;"],
    ["ORA-13609", "Advisor privileges missing",                             "GRANT ADVISOR TO <user>; GRANT EXECUTE ON DBMS_SQLTUNE TO <user>;"],
]
add_table(doc, ["Error Code", "Cause", "Fix"], errors, col_widths=[1.2, 2.4, 3.0])
divider(doc)

# ── Section 8 — Assumptions & Constraints ─────────────────────────────────────
add_heading(doc, "8. Assumptions & Constraints", level=1)
assump = [
    ["Oracle Version",       "12c Release 1 and above (uses FETCH FIRST N ROWS syntax)"],
    ["Python Version",       "Python 3.9 or higher required"],
    ["Read-Only Access",     "The monitoring user requires only SELECT privileges — no DDL/DML access is needed or used"],
    ["Diagnostics Pack",     "AWR-based tabs require Oracle Diagnostics Pack to be licensed on each target database"],
    ["Tuning Pack",          "SQL Tuning Advisor tab requires Oracle Tuning Pack to be licensed on each target database"],
    ["LLM (optional)",       "AI Q&A tab requires Ollama running locally with llama3 (or compatible) model loaded"],
    ["Network",              "Dashboard host must have TCP access to each Oracle listener port"],
    ["Single Instance",      "Dashboard runs on one host; credentials are stored locally in db_configs.json (no shared database)"],
]
add_table(doc, ["Item", "Detail"], assump, col_widths=[2.0, 4.6])
divider(doc)

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run  = footer_para.add_run(
    "Oracle SQL Plan Fluctuation Dashboard  |  Scope & SOP  |  Version 1.0  |  May 2026  |  Confidential"
)
footer_run.font.size      = Pt(8)
footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
footer_run.italic         = True

# ── Save DOCX ─────────────────────────────────────────────────────────────────
doc.save(str(DOCX_OUT))
print(f"[OK] Word document saved: {DOCX_OUT}")

# ── Convert to PDF ─────────────────────────────────────────────────────────────
try:
    from docx2pdf import convert
    convert(str(DOCX_OUT), str(PDF_OUT))
    print(f"[OK] PDF saved: {PDF_OUT}")
except Exception as e:
    print(f"[WARN] PDF conversion failed: {e}")
    print("       You can open the .docx in Word and Save As PDF manually.")
